"""Generador de Memoria Técnica Eléctrica para sistema fotovoltaico.

Documento DOCX firmable por instalador eléctrico autorizado SEC chileno
(clase A/B/C/D según potencia), apto para acompañar la declaración TE-1.
Cumple con las recomendaciones de los Pliegos RIC y la Ley 21.118.

Referencias:
  - Pliego Técnico RIC vigente (SEC Chile)
  - Ley 21.118 (Netbilling) — Generación Distribuida
  - Decreto Supremo N°8 (Reglamento de Seguridad Instalaciones de Consumo)
  - ITG RIC N°9.1/2021 — Instrucción Técnica General
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

from app.services.report_sections import seccion_incluida


ORANGE = RGBColor(0xC5, 0x60, 0x22)
ORANGE_DARK = RGBColor(0x8C, 0x3F, 0x12)
GRAY = RGBColor(0x59, 0x59, 0x59)
BLACK = RGBColor(0x1F, 0x1D, 0x18)


def _cell_shade(cell, hex_color: str):
    """Aplica color de fondo a una celda."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_heading(doc, text, level=1):
    """Heading numerado tipo 1.1, 1.2 con color naranja."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(15); run.font.color.rgb = ORANGE_DARK
    elif level == 2:
        run.font.size = Pt(12); run.font.color.rgb = ORANGE
    else:
        run.font.size = Pt(11); run.font.color.rgb = ORANGE_DARK
    return p


def _p(doc, text, bold=False, italic=False, justify=True, size=10.5, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color: run.font.color.rgb = color
    return p


def _table_kv(doc, rows, col1_cm=5, col2_cm=11):
    """Tabla simple etiqueta : valor."""
    t = doc.add_table(rows=len(rows), cols=2)
    t.autofit = False
    for i, (k, v) in enumerate(rows):
        c1, c2 = t.rows[i].cells
        c1.width = Cm(col1_cm); c2.width = Cm(col2_cm)
        c1.paragraphs[0].clear()
        r = c1.paragraphs[0].add_run(str(k))
        r.font.name = "Calibri"; r.font.size = Pt(10); r.font.bold = True
        r.font.color.rgb = GRAY
        c2.paragraphs[0].clear()
        r2 = c2.paragraphs[0].add_run(str(v))
        r2.font.name = "Calibri"; r2.font.size = Pt(10.5)
        if i % 2 == 1:
            _cell_shade(c1, "FAF6EE"); _cell_shade(c2, "FAF6EE")
    return t


def _table_data(doc, headers, rows, widths_cm=None):
    """Tabla con encabezados y filas de datos."""
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.autofit = False
    if widths_cm:
        for c_idx, w in enumerate(widths_cm):
            for row in t.rows:
                row.cells[c_idx].width = Cm(w)
    # Header
    for c_idx, h in enumerate(headers):
        cell = t.rows[0].cells[c_idx]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(str(h))
        r.font.name = "Calibri"; r.font.size = Pt(9.5); r.font.bold = True
        r.font.color.rgb = ORANGE_DARK
        _cell_shade(cell, "F5E6D8")
    # Rows
    for r_idx, row_data in enumerate(rows, start=1):
        for c_idx, v in enumerate(row_data):
            cell = t.rows[r_idx].cells[c_idx]
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(str(v))
            r.font.name = "Calibri"; r.font.size = Pt(9.5)
            if r_idx % 2 == 0:
                _cell_shade(cell, "FAF6EE")
    return t


def _box_alert(doc, title, text):
    """Caja de alerta/nota destacada."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _cell_shade(cell, "F5E6D8")
    p1 = cell.paragraphs[0]
    p1.clear()
    r = p1.add_run(title)
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = ORANGE_DARK
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    r2 = p2.add_run(text)
    r2.font.name = "Calibri"; r2.font.size = Pt(9.5); r2.font.italic = True; r2.font.color.rgb = GRAY


def _spacer(doc, pts=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)


def generar_word(proyecto: dict, salida: Path | str) -> Path:
    salida = Path(salida)
    fv = proyecto.get("fv", {})
    eco = proyecto.get("eco", {})
    ric = proyecto["ric"]
    s = proyecto["sitio"]
    p_kwp = fv.get("P_kwp", 0)
    n_pan = fv.get("N_paneles", 0)
    es_pmgd = p_kwp > 300
    requiere_te1 = True
    inc = lambda k: seccion_incluida(proyecto, k)

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ===== PORTADA =====
    _spacer(doc, 60)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MEMORIA TÉCNICA")
    r.font.name = "Calibri"; r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = ORANGE
    _spacer(doc, 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sistema Solar Fotovoltaico")
    r.font.name = "Calibri"; r.font.size = Pt(26); r.font.bold = True
    _spacer(doc, 4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Memoria de Cálculo Eléctrico para Declaración TE-1")
    r.font.name = "Calibri"; r.font.size = Pt(12); r.font.italic = True; r.font.color.rgb = GRAY
    _spacer(doc, 40)
    # Caja con datos del proyecto
    t = doc.add_table(rows=4, cols=2)
    for col_idx in range(2):
        for row in t.rows:
            row.cells[col_idx].width = Cm(8)
    datos = [
        ("Proyecto:", proyecto["nombre"]),
        ("Cliente:", proyecto.get("cliente", "—")),
        ("Ubicación:", f"{s['nombre']}, {s['region']}"),
        ("Fecha:", proyecto.get("fecha_creacion", str(date.today()))),
    ]
    for i, (k, v) in enumerate(datos):
        c1, c2 = t.rows[i].cells
        c1.paragraphs[0].clear(); c2.paragraphs[0].clear()
        r1 = c1.paragraphs[0].add_run(k); r1.font.name = "Calibri"; r1.font.size = Pt(11); r1.font.bold = True; r1.font.color.rgb = GRAY
        r2 = c2.paragraphs[0].add_run(v); r2.font.name = "Calibri"; r2.font.size = Pt(12); r2.font.bold = True
        _cell_shade(c1, "FAF6EE"); _cell_shade(c2, "FAF6EE") if i % 2 == 0 else None
    _spacer(doc, 80)
    # Cumplimiento normativo en portada
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Normativa aplicable")
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = ORANGE_DARK
    for txt in [
        "Reglamento de Instalaciones de Consumo (Decreto Supremo N°8)",
        "Pliegos Técnicos RIC vigentes — SEC Chile",
        "Ley 21.118 — Generación Distribuida (Netbilling)",
        "Norma Chilena NCh Elec. 4/2003 — referencia complementaria",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.name = "Calibri"; r.font.size = Pt(10); r.font.color.rgb = GRAY

    doc.add_page_break()

    # ===== 1. ANTECEDENTES GENERALES =====
    _add_heading(doc, "1. Antecedentes Generales", level=1)
    _table_kv(doc, [
        ("Nombre del proyecto", proyecto["nombre"]),
        ("Cliente / titular", proyecto.get("cliente", "—")),
        ("Tipo de proyecto", proyecto["tipo_proyecto"].title()),
        ("Dirección del inmueble", f"{s['nombre']}, {s['region']}, Chile"),
        ("Coordenadas geográficas", f"Lat {s['lat']:.4f}°, Lon {s['lon']:.4f}°"),
        ("Altitud sobre el nivel del mar", f"{s.get('altitud_msnm', 0):.0f} msnm"),
        ("Fecha de elaboración", proyecto.get("fecha_creacion", str(date.today()))),
        ("Régimen de generación", "PMGD (>300 kW)" if es_pmgd else "Netbilling Ley 21.118 (≤300 kW)"),
    ])
    _spacer(doc)
    _p(doc, "El presente documento constituye la memoria técnica de cálculo del sistema solar "
       "fotovoltaico para autoconsumo conectado a red de distribución eléctrica, según la Ley "
       "21.118 de Generación Distribuida y los Pliegos Técnicos RIC de la Superintendencia de "
       "Electricidad y Combustibles (SEC). Se entrega para acompañar la Declaración TE-1 "
       "que efectuará el instalador eléctrico autorizado.")

    # ===== 2. DESCRIPCIÓN DEL PROYECTO =====
    if inc("resumen_ejecutivo"):
        _add_heading(doc, "2. Descripción del Proyecto", level=1)
        _p(doc, f"Se proyecta la instalación de un sistema solar fotovoltaico de "
           f"{p_kwp:.2f} kWp de potencia nominal, compuesto por {n_pan} paneles "
           f"monocristalinos de 550 Wp, montados sobre estructura inclinada con orientación "
           f"al norte y ángulo óptimo determinado para la latitud del sitio. El sistema operará "
           f"en modalidad de autoconsumo con inyección de excedentes a la red de la empresa "
           f"distribuidora local, conforme al marco regulatorio del netbilling chileno.")
        _spacer(doc, 4)
        _p(doc, f"El sistema generará una energía anual estimada de "
           f"{fv.get('generacion_anual_kwh', 0):,.0f} kWh".replace(",", "."), bold=True)
        _p(doc, f"cubriendo aproximadamente el {fv.get('cobertura_real', 0)*100:.0f}% del consumo "
           f"anual estimado del inmueble ({ric['consumo_anual_estimado_kwh']:,.0f} kWh/año), "
           f"y reduciendo las emisiones de gases de efecto invernadero en "
           f"{eco.get('CO2_evitado_anual_kg', 0)/1000:.2f} tCO2 por año, "
           f"considerando el factor de emisión del Sistema Eléctrico Nacional de "
           f"0,200 tCO2/MWh publicado por el Coordinador Eléctrico Nacional.".replace(",", "."))

    # ===== 3. CARACTERIZACIÓN DEL SITIO =====
    _add_heading(doc, "3. Caracterización del Sitio y Recurso Solar", level=1)
    _table_kv(doc, [
        ("Latitud", f"{s['lat']:.4f}°"),
        ("Longitud", f"{s['lon']:.4f}°"),
        ("Altitud", f"{s.get('altitud_msnm', 0):.0f} msnm"),
        ("Inclinación óptima (β)", f"{s['pvgis']['slope']:.0f}°"),
        ("Azimut óptimo (α)", f"{s['pvgis']['azimuth']:.0f}° (orientación norte para hemisferio sur)"),
        ("Irradiación anual plano inclinado H(i)_y", f"{s['pvgis']['H_y']:.0f} kWh/m²/año"),
        ("Generación específica anual E_y", f"{s['pvgis']['E_y']:.0f} kWh/kWp/año"),
        ("Temperatura ambiente promedio anual", f"{s['nasa']['t2m_avg']:.1f} °C"),
        ("Velocidad del viento promedio (10 m)", f"{s['nasa']['wind_avg']:.2f} m/s"),
        ("Fuente de datos solares", "PVGIS v5.3 (Joint Research Centre, Comisión Europea)"),
        ("Fuente de datos meteorológicos", "NASA POWER (Prediction Of Worldwide Energy Resources)"),
    ])

    # ===== 4. CÁLCULO DE CARGAS (RIC) =====
    _add_heading(doc, "4. Memoria de Cálculo de Cargas (RIC)", level=1)
    _p(doc, "El cálculo de carga eléctrica del inmueble se efectúa conforme a las "
       "disposiciones de los Pliegos Técnicos RIC sobre dimensionamiento de circuitos "
       "interiores, considerando la potencia mínima por uso de recinto, las cargas "
       "dedicadas, y los factores de demanda y simultaneidad aplicables al tipo de proyecto.")
    _spacer(doc, 4)
    _add_heading(doc, "4.1 Distribución de recintos", level=2)
    rows = [(c["nombre"], c["uso"], f"{c['area_m2']:.1f}", f"{c['alumbrado_w']:.0f}",
             f"{c['enchufes_w']:.0f}", f"{c['subtotal_w']:.0f}") for c in ric["recintos_carga"]]
    _table_data(doc,
        ["Recinto", "Uso", "Área (m²)", "Alumbrado (W)", "Enchufes (W)", "Subtotal (W)"],
        rows,
        widths_cm=[4.5, 2.5, 2, 2.5, 2.5, 2.5])
    _spacer(doc, 4)
    _add_heading(doc, "4.2 Cargas dedicadas", level=2)
    if ric["cargas_dedicadas"]:
        rows = [(d["nombre"], "Activa" if d.get("activa") else "Inactiva", f"{d['potencia_w']:.0f}")
                for d in ric["cargas_dedicadas"]]
        _table_data(doc, ["Carga", "Estado", "Potencia (W)"], rows, widths_cm=[10, 3, 3.5])
    else:
        _p(doc, "Sin cargas dedicadas declaradas.", italic=True)
    _spacer(doc, 4)
    _add_heading(doc, "4.3 Consolidado y demanda máxima", level=2)
    _table_kv(doc, [
        ("Superficie total construida", f"{ric['area_total_m2']:.1f} m²"),
        ("Subtotal cargas por recinto", f"{ric['subtotal_recintos_w']:,.0f} W".replace(",", ".")),
        ("Subtotal cargas dedicadas", f"{ric['subtotal_dedicadas_w']:,.0f} W".replace(",", ".")),
        ("Carga total instalada", f"{ric['carga_total_instalada_w']:,.0f} W".replace(",", ".")),
        ("Factor de demanda aplicado", f"{ric['factor_demanda']:.2f}"),
        ("Carga diversificada", f"{ric['carga_diversificada_w']:,.0f} W".replace(",", ".")),
        ("Factor de simultaneidad", f"{ric['factor_simultaneidad']:.2f}"),
        ("DEMANDA MÁXIMA DEL EMPALME", f"{ric['demanda_maxima_w']:,.0f} W = {ric['demanda_maxima_w']/1000:.2f} kW".replace(",", ".")),
        ("Corriente nominal del empalme", f"{ric['corriente_nominal_a']:.1f} A"),
        ("Empalme sugerido", ric["tipo_empalme_sugerido"]),
        ("Sistema de conexión", ric["conexion"].replace("_", " ")),
        ("Factor de potencia asumido", f"{ric['factor_potencia']:.2f}"),
    ])

    # ===== 5. DIMENSIONAMIENTO DEL SISTEMA FV =====
    _add_heading(doc, "5. Dimensionamiento del Sistema Fotovoltaico", level=1)
    _table_kv(doc, [
        ("Potencia FV instalada (DC)", f"{p_kwp:.2f} kWp"),
        ("Número de paneles", f"{n_pan} unidades"),
        ("Tecnología", "Silicio monocristalino (PERC / TOPCon)"),
        ("Potencia nominal del panel", "550 Wp (STC)"),
        ("Superficie ocupada (estimada)", f"{fv.get('superficie_m2', 0):.1f} m²"),
        ("Configuración del arreglo", f"Strings calculados según rango MPPT del inversor seleccionado"),
        ("Inversor seleccionado", fv.get("inversor_modelo", "—")),
        ("Potencia AC del inversor", f"{fv.get('inversor_P_AC_kw', 0):.1f} kW"),
        ("Ratio DC/AC", f"1,20"),
        ("Performance Ratio anual estimado", f"{fv.get('PR_anual', 0)*100:.1f} %"),
        ("Factor de planta esperado", f"{fv.get('factor_planta', 0)*100:.1f} %"),
    ])
    _spacer(doc, 4)
    _add_heading(doc, "5.1 Pérdidas técnicas consideradas", level=2)
    if fv.get("perdidas"):
        rows = [(p["nombre"], f"{p['pct']:.2f} %") for p in fv["perdidas"]]
        _table_data(doc, ["Pérdida", "Magnitud"], rows, widths_cm=[10, 4])
        _p(doc, f"Temperatura de celda promedio estimada: {fv.get('T_celda_promedio_c', 25):.1f} °C "
           f"(modelo Sandia con altura solar local y corrección por altitud).", italic=True)
    _spacer(doc, 4)
    _add_heading(doc, "5.2 Producción esperada", level=2)
    _table_kv(doc, [
        ("Generación anual estimada", f"{fv.get('generacion_anual_kwh', 0):,.0f} kWh".replace(",", ".")),
        ("Cobertura solar del consumo", f"{fv.get('cobertura_real', 0)*100:.1f} %"),
        ("Energía autoconsumida (estimada)", f"{fv.get('autoconsumo_kwh', 0):,.0f} kWh/año".replace(",", ".")),
        ("Energía inyectada a red (estimada)", f"{fv.get('inyeccion_kwh', 0):,.0f} kWh/año".replace(",", ".")),
        ("Energía comprada a red (residual)", f"{fv.get('compra_red_kwh', 0):,.0f} kWh/año".replace(",", ".")),
    ])

    # ===== 6. PROTECCIONES Y SEGURIDAD =====
    _add_heading(doc, "6. Protecciones Eléctricas y Seguridad", level=1)
    _add_heading(doc, "6.1 Lado DC", level=2)
    _p(doc, "Cada string de paneles deberá estar protegido mediante:")
    for txt in [
        "Fusibles tipo gPV en serie con cada string, dimensionados según Isc del panel × 1,56",
        "Interruptor seccionador DC bipolar con capacidad de corte bajo carga",
        "Descargador de sobretensiones (SPD) tipo II clase II en el string box",
        "Conductor de protección (PE) continuo desde estructura hasta puesta a tierra",
    ]:
        p = doc.add_paragraph(style="List Bullet"); p.add_run(txt).font.size = Pt(10.5)
    _add_heading(doc, "6.2 Lado AC", level=2)
    _p(doc, "A la salida del inversor se instalarán:")
    for txt in [
        "Interruptor automático magnetotérmico, In dimensionada según corriente nominal del inversor",
        "Interruptor diferencial superinmunizado, sensibilidad 30 mA tipo A o B según fabricante del inversor",
        "Descargador de sobretensiones (SPD) tipo II clase II",
        "Medidor bidireccional para netbilling, suministrado y/o aprobado por la empresa distribuidora",
    ]:
        p = doc.add_paragraph(style="List Bullet"); p.add_run(txt).font.size = Pt(10.5)
    _add_heading(doc, "6.3 Puesta a tierra", level=2)
    _p(doc, "El sistema fotovoltaico se conectará a la malla de tierra existente del inmueble. "
       "Si no existe malla adecuada, se construirá una nueva con resistencia de puesta a tierra "
       "menor o igual a 25 Ω, medida con telurómetro calibrado en el momento de la puesta en "
       "servicio. La estructura metálica de soporte de los paneles y todas las partes metálicas "
       "accesibles deben equipotencializarse con conductor de cobre desnudo de sección mínima "
       "según la corriente de cortocircuito esperada.")

    # ===== 7. CUMPLIMIENTO NORMATIVO =====
    _add_heading(doc, "7. Cumplimiento Normativo", level=1)
    _table_kv(doc, [
        ("Categoría netbilling", "BT1 ≤ 20 kW (compra obligatoria de excedentes)" if p_kwp <= 20
                                  else ("Netbilling estándar (20–300 kW)" if not es_pmgd else "PMGD (>300 kW)")),
        ("Ley aplicable", "Ley 21.118 — Generación Distribuida"),
        ("Norma técnica eléctrica", "Pliegos Técnicos RIC vigentes (SEC Chile)"),
        ("Norma anti-isla", "IEEE 1547 / VDE-AR-N 4105 (función nativa del inversor)"),
        ("Homologación del inversor", "Equipos homologados por SEC con número de certificación vigente"),
        ("Declaración SEC requerida", "TE-1 (instalador autorizado clase A/B/C/D según potencia y tensión)"),
        ("Distribuidora", "Trámite de conexión y aprobación según protocolo de la empresa local"),
    ])
    if p_kwp >= 20:
        _spacer(doc, 4)
        _box_alert(doc, "Memoria explicativa obligatoria",
                   f"La instalación supera los 20 kW de potencia, por tanto la memoria explicativa "
                   f"es OBLIGATORIA para el trámite TE-1 ante SEC. Este documento cumple con esa "
                   f"exigencia. La potencia proyectada es de {p_kwp:.2f} kWp.")
    if es_pmgd:
        _spacer(doc, 4)
        _box_alert(doc, "Régimen PMGD",
                   f"La potencia ({p_kwp:.2f} kWp) excede los 300 kW del netbilling. El proyecto "
                   f"debe tramitarse como Pequeño Medio de Generación Distribuida (PMGD), lo que "
                   f"implica estudios adicionales (calidad de servicio, EIA si > 3 MW, etc.).")

    # ===== 8. OPERACIÓN Y MANTENIMIENTO =====
    _add_heading(doc, "8. Operación y Mantenimiento", level=1)
    _p(doc, "Para garantizar la vida útil del sistema (mínimo 25 años para los módulos) se "
       "recomienda el siguiente programa de mantenimiento preventivo:")
    for txt in [
        "Limpieza de paneles cada 3 a 6 meses según el nivel de suciedad de la zona (más frecuente en zona norte / agroindustria)",
        "Inspección visual de cableado, conectores MC4 y sellos de paso de cubierta — anualmente",
        "Termografía con cámara IR para detectar puntos calientes y mismatch — anualmente",
        "Verificación del valor de aislación DC y AC con megóhmetro — anualmente",
        "Medición de la puesta a tierra con telurómetro — anualmente",
        "Reposición del inversor estimada al año 12 (60% del costo original)",
        "Reposición de baterías (si aplica) al año 10 para tecnología LFP",
    ]:
        p = doc.add_paragraph(style="List Bullet"); p.add_run(txt).font.size = Pt(10.5)

    # ===== 9. ANEXOS =====
    _add_heading(doc, "9. Anexos al Proyecto", level=1)
    _p(doc, "Deben acompañar esta memoria los siguientes documentos:")
    for txt in [
        "Plano de implantación del arreglo fotovoltaico (escala 1:50 o 1:100)",
        "Diagrama unilineal eléctrico del sistema, con identificación de protecciones",
        "Fichas técnicas (datasheet) de los módulos y del inversor seleccionado",
        "Certificados de homologación SEC del inversor",
        "Cálculos de caída de tensión DC y AC",
        "Protocolo de ensayos: aislación, continuidad de PE, medición de tierra, prueba de funcionamiento",
        "Foto-fija de la instalación terminada antes de la puesta en servicio",
    ]:
        p = doc.add_paragraph(style="List Bullet"); p.add_run(txt).font.size = Pt(10.5)

    # ===== 10. FIRMA =====
    doc.add_page_break()
    _add_heading(doc, "10. Firma del Instalador Eléctrico Autorizado", level=1)
    _p(doc, "El instalador eléctrico abajo firmante declara bajo juramento que el proyecto "
       "fotovoltaico descrito en esta memoria cumple con las disposiciones del Reglamento "
       "de Seguridad de las Instalaciones de Consumo de Energía Eléctrica (Decreto N°8) "
       "y los Pliegos Técnicos RIC vigentes, y que la ejecución se realizará bajo su "
       "directa responsabilidad.")
    _spacer(doc, 40)
    t = doc.add_table(rows=6, cols=2)
    for col in range(2):
        for row in t.rows: row.cells[col].width = Cm(8)
    firma = [
        ("Nombre completo:", "_____________________________________"),
        ("RUT:", "_____________________________________"),
        ("Clase de instalador:", "  A  □    B  □    C  □    D  □"),
        ("N° de Licencia SEC:", "_____________________________________"),
        ("Teléfono / contacto:", "_____________________________________"),
        ("Lugar y fecha:", "_____________________________________"),
    ]
    for i, (k, v) in enumerate(firma):
        c1, c2 = t.rows[i].cells
        c1.paragraphs[0].clear(); c2.paragraphs[0].clear()
        r1 = c1.paragraphs[0].add_run(k); r1.font.name = "Calibri"; r1.font.size = Pt(11); r1.font.bold = True; r1.font.color.rgb = GRAY
        r2 = c2.paragraphs[0].add_run(v); r2.font.name = "Calibri"; r2.font.size = Pt(11)
    _spacer(doc, 60)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("________________________________________")
    r.font.name = "Calibri"; r.font.size = Pt(11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Firma y timbre del instalador eléctrico autorizado")
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.italic = True; r.font.color.rgb = GRAY

    # Pie de portada-final
    _spacer(doc, 30)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Documento generado por FV Chile · v0.6 · {date.today().isoformat()}")
    r.font.name = "Calibri"; r.font.size = Pt(8); r.font.color.rgb = GRAY

    doc.save(salida)
    return salida
